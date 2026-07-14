"""
REVISED draft questionnaire (v2) — incorporates Lamia's review feedback:
  - Adaptive flow (choose problem → only relevant follow-ups) — 6-12 questions
  - Wording edits (প্রায়, গত ১২ মাসে, অনুমতি ছাড়া, ইচ্ছাকৃতভাবে, ...)
  - Medical neglect as a separate question
  - Gender + District (for real-time dashboard)
  - Abuser: family? → then relationship option buttons
  - Risk: "এখনই নিরাপদ জায়গায় যাওয়ার প্রয়োজন?"
  - Result page: + Confidence Score + AI explanation
  - UI: 🔊 শুনুন on every page, progress bar

Output: C:\\Users\\ACER\\Downloads\\question_bank_draft_v2.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"C:\Users\ACER\Downloads\question_bank_draft_v2.docx")
BN = "Nirmala UI"
BLUE = (0x1a, 0x3c, 0x6e)
GREY = (0x80, 0x80, 0x80)
ORANGE = (0xB0, 0x5A, 0x00)
GREEN = (0x1E, 0x7A, 0x3C)


def r(p, text, size=10.5, bold=False, color=None):
    run = p.add_run(text)
    run.font.name = BN
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def para(doc, text="", size=10.5, bold=False, color=None, style=None, center=False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r(p, text, size, bold, color)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r(c.paragraphs[0], h, 10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r(cells[i].paragraphs[0], str(v), 10)
    return t


def build():
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    para(doc, "প্রশ্নপত্র খসড়া v2 — Elder Abuse Detection App", 16, True, BLUE, center=True)
    para(doc, "Adaptive Question Flow (Lamia-র রিভিউ অনুযায়ী সংশোধিত)", 11, color=GREY, center=True)
    para(doc)

    # ── What changed ─────────────────────────────────────────────────────────
    para(doc, "v1 থেকে যা পরিবর্তন হয়েছে", 13, True, GREEN)
    for line in [
        "❌ সবাইকে ১৭টা প্রশ্ন নয় → ✅ Adaptive: সমস্যা বেছে নিলে শুধু সেটার প্রশ্ন (৬–১২টি)",
        "✅ বয়সে \"প্রায়\" যোগ (বৃদ্ধরা সঠিক বয়স জানেন না)",
        "✅ শারীরিক: \"গত ১২ মাসে\" সময়সীমা + ধাক্কা/চড়/লাথি যোগ",
        "✅ আর্থিক: \"আপনার অনুমতি ছাড়া\" যোগ",
        "✅ অবহেলা: \"ইচ্ছাকৃতভাবে\" + পানি যোগ; চিকিৎসা-অবহেলা আলাদা প্রশ্ন",
        "✅ পরিত্যাগ: \"দীর্ঘ সময় একা ফেলে রাখা\" যোগ",
        "✅ নির্যাতনকারী: আগে \"পরিবারের সদস্য?\" → তারপর সম্পর্ক বাছাই (button)",
        "✅ ঝুঁকি: \"এখনই নিরাপদ জায়গায় যাওয়ার প্রয়োজন আছে?\" যোগ",
        "🆕 Gender + জেলা (Dashboard-এর real-time analytics-এর জন্য)",
        "🆕 ফলাফলে: Confidence Score + AI-এর ব্যাখ্যা",
        "🆕 প্রতিটি screen-এ 🔊 \"শুনুন\" button + Progress bar",
    ]:
        para(doc, line, 10.5, style="List Bullet")
    para(doc)

    # ── Flow ─────────────────────────────────────────────────────────────────
    doc.add_page_break()
    para(doc, "সম্পূর্ণ Flow", 13, True, BLUE)
    for line in [
        "১. কে রিপোর্ট করছে?      [ আমি নিজে ]  [ অন্যের পক্ষে ]",
        "২. বয়স প্রায় ৬০+?        [ হ্যাঁ ]  [ না ]",
        "৩. ভুক্তভোগী পুরুষ/মহিলা?  [ পুরুষ ]  [ মহিলা ]          ← Dashboard",
        "৪. জেলা বাছাই             [ ৬৪ জেলার তালিকা ]            ← Dashboard",
        "",
        "৫. ⭐ সমস্যা বেছে নিন (একাধিক বাছা যাবে — বড় icon button):",
        "        👊 মারধর / আঘাত",
        "        🗣️ গালিগালাজ / হুমকি",
        "        💰 টাকা / সম্পত্তি জোর করে নেওয়া",
        "        🍚 খাবার / ওষুধ দেয় না",
        "        🏠 বাড়ি থেকে বের / একা ফেলে রাখা",
        "        ⚫ মারা গেছেন   (শুধু \"অন্যের পক্ষে\" হলে)",
        "",
        "৬. শুধু নির্বাচিত সমস্যার follow-up প্রশ্ন (প্রতিটির ১–২টি)",
        "",
        "৭. নির্যাতনকারী কে?  → পরিবারের সদস্য? → সম্পর্ক বাছাই",
        "৮. ঝুঁকি (২টি প্রশ্ন)",
        "",
        "৯. AI বিশ্লেষণ → ফলাফল (🔊 পড়ে শোনাবে)",
    ]:
        para(doc, line, 10.5)
    para(doc)
    para(doc, "→ একটি সমস্যা বাছলে মোট প্রশ্ন: ~১০–১১টি   |   দুইটি সমস্যা: ~১২–১৩টি",
         10.5, True, GREEN)

    # ── Sections ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    para(doc, "ধাপ ১–৪ — শুরু (সবার জন্য)", 13, True, BLUE)
    table(doc,
          ["#", "প্রশ্ন (বাংলা)", "উত্তরের ধরন", "কী কাজে লাগে"],
          [
              ["১", "আপনি কি নিজে এই সমস্যার শিকার?", "[আমি নিজে] [অন্যের পক্ষে]",
               "reporter type (murder প্রশ্ন দেখাবে কি না)"],
              ["২", "ভুক্তভোগীর বয়স কি প্রায় ৬০ বছর বা তার বেশি?", "[হ্যাঁ] [না]",
               "elder confirm"],
              ["৩", "ভুক্তভোগী কি পুরুষ না মহিলা?", "[পুরুষ] [মহিলা]",
               "Dashboard — gender analytics"],
              ["৪", "আপনি কোন জেলায় থাকেন?", "৬৪ জেলার তালিকা",
               "Dashboard — district map"],
          ])
    para(doc)

    para(doc, "ধাপ ৫ — সমস্যা বেছে নিন (একাধিক বাছা যাবে)", 13, True, BLUE)
    para(doc, "বড় icon button — পড়তে না পারলেও ছবি দেখে চেনা যাবে; TTS পড়ে শোনাবে।",
         9.5, color=GREY)
    table(doc,
          ["Icon", "লেখা (বাংলা)", "Category"],
          [
              ["👊", "মারধর / আঘাত", "physical"],
              ["🗣️", "গালিগালাজ / হুমকি", "verbal"],
              ["💰", "টাকা / সম্পত্তি জোর করে নেওয়া", "financial"],
              ["🍚", "খাবার / ওষুধ দেয় না", "neglect"],
              ["🏠", "বাড়ি থেকে বের / একা ফেলে রাখা", "abandonment"],
              ["⚫", "মারা গেছেন  (শুধু অন্যের পক্ষে হলে)", "murder"],
          ])
    para(doc)

    # ── Follow-ups ───────────────────────────────────────────────────────────
    doc.add_page_break()
    para(doc, "ধাপ ৬ — Follow-up প্রশ্ন (শুধু যেটা বাছা হয়েছে)", 13, True, BLUE)

    followups = [
        ("👊 শারীরিক নির্যাতন (physical) — severity 4 · দণ্ডবিধি §323, §324; PMA §3", [
            ["P1", "গত ১২ মাসে কি আপনাকে মারধর, ধাক্কা, চড় বা লাথি দেওয়া হয়েছে?", "[হ্যাঁ] [না]"],
            ["P2", "আঘাতের কারণে কি কখনো ডাক্তার বা হাসপাতালে যেতে হয়েছে?", "[হ্যাঁ] [না]"],
        ]),
        ("🗣️ মৌখিক / মানসিক নির্যাতন (verbal) — severity 1 · দণ্ডবিধি §506", [
            ["V1", "আপনাকে কি গালিগালাজ, অপমান, ভয় দেখানো বা হুমকি দেওয়া হয়?", "[হ্যাঁ] [না]"],
            ["V2", "আপনাকে কি একঘরে করে রাখা হয় বা \"বোঝা\" মনে করা হয়?", "[হ্যাঁ] [না]"],
        ]),
        ("💰 আর্থিক নির্যাতন (financial) — severity 2 · দণ্ডবিধি §406, §420", [
            ["F1", "আপনার অনুমতি ছাড়া কি কেউ আপনার টাকা, পেনশন বা বয়স্ক ভাতা নিয়ে নেয়?", "[হ্যাঁ] [না]"],
            ["F2", "আপনার অনুমতি ছাড়া কি জমি বা সম্পত্তি লিখে নেওয়া হয়েছে?", "[হ্যাঁ] [না]"],
        ]),
        ("🍚 অবহেলা (neglect) — severity 2 · PMA §4", [
            ["N1", "পরিবার কি ইচ্ছাকৃতভাবে আপনাকে খাবার, পানি বা যত্ন থেকে বঞ্চিত করে?", "[হ্যাঁ] [না]"],
            ["N2", "অসুস্থ হলে পরিবার কি ইচ্ছাকৃতভাবে চিকিৎসা বা ওষুধ দেয় না?", "[হ্যাঁ] [না]"],
        ]),
        ("🏠 পরিত্যাগ (abandonment) — severity 3 · PMA §3, §4", [
            ["A1", "আপনাকে কি বাড়ি থেকে বের করে দেওয়া হয়েছে?", "[হ্যাঁ] [না]"],
            ["A2", "আপনাকে কি দীর্ঘ সময় একা ফেলে রাখা হয় বা পরিবার খোঁজ নেয় না?", "[হ্যাঁ] [না]"],
        ]),
        ("⚫ হত্যা (murder) — severity 5 · দণ্ডবিধি §302, §304 · শুধু \"অন্যের পক্ষে\"", [
            ["M1", "নির্যাতনের কারণে কি ভুক্তভোগী মারা গেছেন?", "[হ্যাঁ] [না]  → জরুরি: ৯৯৯ + থানা"],
        ]),
    ]
    for title, rows in followups:
        para(doc, title, 11.5, True)
        table(doc, ["#", "প্রশ্ন (বাংলা)", "উত্তর"], rows)
        para(doc)

    # ── Abuser + Risk ────────────────────────────────────────────────────────
    doc.add_page_break()
    para(doc, "ধাপ ৭ — নির্যাতনকারী কে? (Trust Blind Spot)", 13, True, BLUE)
    para(doc, "Dataset finding: ৪১.৮% নির্যাতনকারী পরিবারের সদস্য।", 9.5, color=GREY)
    table(doc,
          ["#", "প্রশ্ন / পর্দা", "উত্তরের ধরন"],
          [
              ["AB1", "নির্যাতনকারী কি পরিবারের সদস্য?", "[হ্যাঁ] [না]"],
              ["AB2a", "সম্পর্ক বেছে নিন  (যদি AB1 = হ্যাঁ)",
               "[ছেলে] [মেয়ে] [পুত্রবধূ] [জামাই] [নাতি-নাতনি] [ভাই-বোন] [অন্য আত্মীয়]"],
              ["AB2b", "কে?  (যদি AB1 = না)", "[প্রতিবেশী] [অপরিচিত] [অন্য]"],
          ])
    para(doc)

    para(doc, "ধাপ ৮ — ঝুঁকি / জরুরি অবস্থা", 13, True, BLUE)
    table(doc,
          ["#", "প্রশ্ন (বাংলা)", "উত্তর", "প্রভাব"],
          [
              ["R1", "নির্যাতন কি এখনো চলছে?", "[হ্যাঁ] [না]", "চলমান → risk বাড়ে"],
              ["R2", "আপনার কি এখনই নিরাপদ জায়গায় যাওয়ার প্রয়োজন আছে?", "[হ্যাঁ] [না]",
               "হ্যাঁ → জরুরি SOS: ৯৯৯"],
          ])
    para(doc)

    # ── Result page ──────────────────────────────────────────────────────────
    doc.add_page_break()
    para(doc, "ধাপ ৯ — ফলাফল পর্দা (Result Page)", 13, True, BLUE)
    table(doc,
          ["যা দেখাবে", "উদাহরণ"],
          [
              ["নির্যাতনের ধরন (Abuse Type)", "শারীরিক নির্যাতন + অবহেলা"],
              ["ঝুঁকির মাত্রা (Risk Level)", "উচ্চ"],
              ["Confidence Score  🆕", "৮৭%"],
              ["লঙ্ঘিত আইন (Violated Laws)", "PMA 2013 §3, §4 · দণ্ডবিধি §323"],
              ["AI-এর ব্যাখ্যা  🆕",
               "\"আপনি বলেছেন গত ১২ মাসে মারধর হয়েছে এবং হাসপাতালে যেতে হয়েছে — "
               "তাই এটি শারীরিক নির্যাতন। খাবার থেকে বঞ্চিত করা হয় — তাই অবহেলাও।\""],
              ["করণীয় (Recommendations)", "থানায় FIR · UNO-তে অভিযোগ · NLASO ১৬৪৩০"],
              ["জরুরি নম্বর (Emergency)", "৯৯৯ · ১৬৪৩০ · বিশ্বস্ত ব্যক্তিকে জানান"],
          ])
    para(doc)

    para(doc, "প্রতিটি পর্দায় (UI)", 12, True, GREEN)
    for line in [
        "🔊 \"শুনুন\" button — প্রশ্ন/ফলাফল আবার পড়ে শোনাবে",
        "Progress bar — যেমন: \"প্রশ্ন ৪ / ১০\"",
        "বড় font, বড় button, রঙ-ভিত্তিক (সবুজ = হ্যাঁ, লাল = না)",
    ]:
        para(doc, line, 10.5, style="List Bullet")
    para(doc)

    # ── Dashboard integration ────────────────────────────────────────────────
    para(doc, "Dashboard Integration (real-time)", 13, True, BLUE)
    para(doc, "প্রতিটি সম্পূর্ণ assessment → Firebase Firestore-এ সংরক্ষণ → "
              "Admin Dashboard স্বয়ংক্রিয়ভাবে আপডেট", 10.5)
    table(doc,
          ["সংরক্ষিত হবে", "কোথা থেকে"],
          [
              ["Abuse Type", "AI বিশ্লেষণ"],
              ["Risk Level", "AI বিশ্লেষণ"],
              ["District", "প্রশ্ন ৪"],
              ["Gender", "প্রশ্ন ৩"],
              ["Time (timestamp)", "স্বয়ংক্রিয়"],
              ["Abuser relation", "ধাপ ৭"],
          ])
    para(doc)
    para(doc, "⚠️ কোনো নাম/পরিচয় সংরক্ষণ করা হবে না — শুধু anonymized পরিসংখ্যান।",
         10, True, ORANGE)
    para(doc)

    # ── Review box ───────────────────────────────────────────────────────────
    doc.add_page_break()
    para(doc, "রিভিউ — আপনার মতামত এখানে লিখুন", 13, True, ORANGE)
    for line in [
        "কোন প্রশ্নের শব্দ পরিবর্তন করতে চান?   →",
        "",
        "কোন প্রশ্ন যোগ করতে চান?   →",
        "",
        "কোন প্রশ্ন বাদ দিতে চান?   →",
        "",
        "\"সমস্যা বেছে নিন\" পর্দার icon/লেখা ঠিক আছে?   →",
        "",
        "অন্য মতামত:   →",
        "",
    ]:
        para(doc, line, 10.5)

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Size : {OUT.stat().st_size/1024:.1f} KB")


if __name__ == "__main__":
    build()
